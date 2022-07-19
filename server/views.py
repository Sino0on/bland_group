import time

from django.contrib.auth import login, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.shortcuts import render, redirect

from .filters import ProductFilter
from .models import *
from .forms import *
import requests
from django.shortcuts import render, redirect, get_object_or_404
from django.views.decorators.http import require_POST
from .cart import Cart
from .forms import CartAddProductForm

from docx import Document as Doc
from django.core.files.base import ContentFile, File


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        # print(inline)
        for item in inline:
            # print(item)
            if key in item.text:
                # print(key)
                item.text = item.text.replace(key, value)
            else:
                print(item.text)


@require_POST
def cart_add(request, product_id):
    cart = Cart(request)
    product = get_object_or_404(Product, id=product_id)
    form = CartAddProductForm(request.POST)
    if form.is_valid():
        cd = form.cleaned_data
        cart.add(product=product,
                 quantity=cd['quantity'],
                 update_quantity=cd['update'])
    return redirect('cart_detail')


def cart_remove(request, product_id):
    cart = Cart(request)
    product = get_object_or_404(Product, id=product_id)
    cart.remove(product)
    return redirect('cart_detail')


def cart_detail(request):
    cart = Cart(request)
    return render(request, 'cart_detail.html', {'cart': cart})


def register(request):  # функция регистрации
    print('da')
    if request.method == 'POST':  # Проверка запроса на пост
        form = UserCreationForm(request.POST)  # присваиваем форму для данных
        print('POST')
        if form.is_valid():  # Проверка на валидность
            form.save()  # Сохранение в базу
            print('VAlid')
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = authenticate(username=username, password=password)  # авторизация юзера
            print('hz')
            login(request, user)  # авторизация юзера
            return redirect('/')  # переадресация на главную страничку
        else:
            print(form.errors)
    else:  # если это запрос не пост
        form = UserCreationForm()  # Присваивание форму

    context = {'form': form, 'cart': Cart(request)}  # контекст для передачи данных для шаблона
    return render(request, 'register.html', context)


def index(request):
    if request.method == 'POST':
        print('post')
        form = FeedbackCreateForm(request.POST)
        if form.is_valid():
            print('valid')
            das = form.save()
            payload = {
                "text": f"Имя - {das.name}\nПочта - {das.email}\nТекст - {das.text}\nВремя - {das.date}",
                "chat_id": '572219148',
            }
            response = requests.post(
                'https://api.telegram.org/bot5324134927:AAEecQYFSOAnpX--_awWaACuUQxGofyrk2w/sendMessage', json=payload)
            print('ok')
            return redirect('/')
        else:
            form = form
    else:
        form = FeedbackCreateForm
    categories = Category.objects.all()
    form2 = AuthenticationForm
    return render(request, 'index.html', {'form': form, 'form2': form2, 'cart': Cart(request), 'categories': categories, 'subcategories': Subcategory.objects.all()})


def product_list(request):
    products = Product.objects.all()
    filter = ProductFilter(request.GET, queryset=products)
    products = filter.qs
    return render(request, 'products.html', {'products': products, 'cart_product_form': CartAddProductForm, 'productfilter': filter})


def postdelete(request, id):
    das = Product.objects.get(id=id)
    das.delete()
    return redirect('/')


def news_list(request):
    news = News.objects.all()
    return render(request, 'news_list.html', {'news': news, 'cart': Cart(request)})


def product_create(request):
    if request.method == 'POST':
        print('post')
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            print('valid')

            das = form.save()
            total = 0
            product = Product.objects.get(id=das.id)
            images = request.FILES.getlist('files')
            for image in images:
                Image_product.objects.create(product=product, image=image)
            payload = {
                "text": f"Только что создали продукт под названием {das.title}\nМини-название - {das.mini_title}\nКатегория - {das.category}\nC {total} изображениями",
                "chat_id": '795677145',
            }
            response = requests.post(
                'https://api.telegram.org/bot5324134927:AAEecQYFSOAnpX--_awWaACuUQxGofyrk2w/sendMessage', json=payload)
            print('ok')
            return redirect('/')
        else:
            form = form
    else:
        form = ProductForm
    return render(request, 'product_create.html', {'from': form, 'cart': Cart(request)})


def product_detail(request, pk):
    product = Product.objects.get(id=pk)
    return render(request, 'product_detail.html', {'product': product})


def components(request):
    return render(request, 'components.html', {'cart': Cart(request)})


def orders(request):
    return render(request, 'orders.html', {'orders': OrderUr.objects.all()})


def test(request):
    return render(request, 'html.html', {'cart': Cart(request)})


def orderfiz_create(request):
    if request.method == 'POST':
        form = OrderFizCreateForm(request.POST)
        if form.is_valid():
            das = form.save(commit=False)
            das.user = request.user

            variables = {
                "${NAME}": str(das.first_name) + ' ' + str(das.last_name),
                "${INN}": 'ИНН ' + str(das.series),
                "${SERI}": f'паспорт {das.series_type} выдан {das.series}',
                "${DATE}": 'дата выдачи '+str(das.date_of_issue),
                "${ADRES}": 'зарегистрированный по адресу: '+das.adress,
                "${EMAIL}": 'Эл.почта: '+str(das.email),
                "${PHONE}": 'тел.: '+str(das.phone),
            }

            template_document = Doc('/Users/user/PycharmProjects/blend/blend/server/dogovor_fiz.docx')
            for i in variables:
                for p in template_document.paragraphs:
                    if p.text.find(i) >= 0:
                        print('yes')
                        print(i)
                        p.text = p.text.replace(i, variables[i])
            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                print(paragraph.text)
                                # for i in paragraph.runs:
                                #     print(i)
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            # for table in template_document.tables:
            #     for col in table:
            #         for cell in col.cells:
            #             for paragraph in cell.paragraphs:
            #                 print(paragraph.runs)
            records = (
                (3, '101', 'Spam'),
                (7, '422', 'Eggs'),
                (4, '631', 'Spam, spam, eggs, and spam')
            )
            cart = Cart(request)
            for i in cart:
                print(i)
            table = template_document.add_table(rows=1, cols=6,  style="Table Grid")
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Название'
            hdr_cells[2].text = 'Мини-Название'
            hdr_cells[3].text = 'Категория'
            hdr_cells[4].text = 'Колличество'
            hdr_cells[5].text = 'Цена'
            for i in cart:
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = str(i['product'].title)
                row_cells[2].text = i['product'].mini_title
                row_cells[3].text = i['product'].category.title
                row_cells[4].text = str(i['quantity'])
                row_cells[5].text = str(i['total_price'])
            row_cells = table.add_row().cells
            row_cells[0].text = ''
            row_cells[1].text = ''
            row_cells[2].text = ''
            row_cells[3].text = ''
            row_cells[4].text = ''
            row_cells[5].text = str(cart.get_total_price())
            docume = template_document.save(f'media/{das.last_name}-{time.time()}.docx')
            das.document = f'{das.last_name}-{time.time()}.docx'
            # print(docume)
            # print(type(docume))
            # with open(docume, 'rb') as f:
            #     das.product.save('dogobor.docx', File(f))
            das.save()

            return redirect('/')
        else:
            print(form.errors)
            return render(request, 'order_create_fiz.html', {'cart': Cart(request), 'form': form})
    else:
        form = OrderFizCreateForm()
    return render(request, 'order_create_fiz.html', {'cart': Cart(request), 'form': form})


def orderur_create(request):
    if request.method == 'POST':
        form = OrderUrCreateForm(request.POST)
        if form.is_valid():
            das = form.save(commit=False)
            das.user = request.user

            variables = {
                "${NAME}": str(das.name),
                "${INN}": str(das.inn),
                "${OKPO}": f'{das.okpo}',
                "${BIK}": str(das.bik),
                "${RC}": f'{das.r_c}',
                "${BANK}": f'{das.bank}',
                "${EMAIL}": str(das.email),
                "${PHONE}": str(das.phone),
                "${DIRECTOR}": str(das.director),
            }

            template_document = Doc('/Users/user/PycharmProjects/blend/blend/server/dogovor_ur.docx')
            for i in variables:
                for p in template_document.paragraphs:
                    if p.text.find(i) >= 0:
                        # print('yes')
                        # print(i)
                        p.text = p.text.replace(i, variables[i])
            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                # print(paragraph.text)
                                # for i in paragraph.runs:
                                #     print(i)
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            # for table in template_document.tables:
            #     for col in table:
            #         for cell in col.cells:
            #             for paragraph in cell.paragraphs:
            #                 print(paragraph.runs)
            records = (
                (3, '101', 'Spam'),
                (7, '422', 'Eggs'),
                (4, '631', 'Spam, spam, eggs, and spam')
            )
            cart = Cart(request)
            for i in cart:
                print(i)
            table = template_document.add_table(rows=1, cols=6,  style="Table Grid")
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Название'
            hdr_cells[2].text = 'Мини-Название'
            hdr_cells[3].text = 'Категория'
            hdr_cells[4].text = 'Колличество'
            hdr_cells[5].text = 'Цена'
            for i in cart:
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = str(i['product'].title)
                row_cells[2].text = i['product'].mini_title
                row_cells[3].text = i['product'].category.title
                row_cells[4].text = str(i['quantity'])
                row_cells[5].text = str(i['total_price'])
            row_cells = table.add_row().cells
            row_cells[0].text = ''
            row_cells[1].text = ''
            row_cells[2].text = ''
            row_cells[3].text = ''
            row_cells[4].text = ''
            row_cells[5].text = str(cart.get_total_price()+cart.get_total_price()/100*16)
            docume = template_document.save(f'media/{das.name}-{time.time()}.docx')
            das.document = f'{das.name}-{time.time()}.docx'
            cart.clear()
            return redirect('/')
        else:
            print(form.errors)
            return render(request, 'order_create_ur.html', {'cart': Cart(request), 'form': form})
    else:
        form = OrderUrCreateForm()
    return render(request, 'order_create_ur.html', {'cart': Cart(request), 'form': form})


# def generate_docx(loan_id):
#     documents = Document.objects.filter(is_active=True)
#     template_file_path = documents[0].document
#
#     order = OrderFiz.objects.get(pk=loan_id)
#
#     date = datetime.date.today()
#     # Data
#     variables = {
#         "${EMPLOEE_NAME}": order.full_name,
#         "${EMPLOEE_TITLE}": order.position,
#         "${EMPLOEE_ID}": order.id,
#         "${EMPLOEE_ADDRESS}": order.address,
#         "${EMPLOEE_PHONE}": "example_phone_number",
#         "${START_DATE}": date,
#     }
#
#     template_document = Doc(template_file_path)
#     for i in variables:
#         for p in template_document.paragraphs:
#             if p.text.find(i) >= 0:
#                 p.text = p.text.replace(i, variables[i])
#     for variable_key, variable_value in variables.items():
#         for paragraph in template_document.paragraphs:
#             replace_text_in_paragraph(paragraph, variable_key, variable_value)
#
#         for table in template_document.tables:
#             for col in table.columns:
#                 for cell in col.cells:
#                     for paragraph in cell.paragraphs:
#                         replace_text_in_paragraph(paragraph, variable_key, variable_value)
#     return template_document
